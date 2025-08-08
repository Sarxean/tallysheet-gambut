# file: tallysheet_docx.py
from fastapi import FastAPI, Form, BackgroundTasks, File, UploadFile
from fastapi.responses import FileResponse, JSONResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from io import BytesIO
import tempfile
import uuid
import os
from typing import Optional, Dict, Any, List, Tuple, Union

app = FastAPI()

# ----------------------- Helpers -----------------------
def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def format_input_text(value: Optional[str], placeholder_length: int = 6) -> str:
    if value is None or str(value).strip() == "":
        return "_" * placeholder_length
    return str(value)


def format_checkbox_as_X(value: Optional[str]) -> str:
    if value is None or str(value).strip() == "":
        return "_"
    return "X"


def set_run_style(run, name: str = "Cambria", size_pt: int = 11, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def _insert_image_in_cell(cell, image_bytes_or_path: Optional[Union[bytes, str]], max_width_inches: float = 2.8):
    """
    Insert image centered in a table cell. Accepts bytes (image content) or filesystem path (str).
    If no image, insert placeholder text.
    """
    # Clear existing paragraphs (keep first)
    for p in cell.paragraphs:
        p.clear()

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not image_bytes_or_path:
        run = p.add_run("(no image)")
        set_run_style(run, size_pt=10, italic=True)
        return

    try:
        if isinstance(image_bytes_or_path, (bytes, bytearray)):
            stream = BytesIO(image_bytes_or_path)
            p.add_run().add_picture(stream, width=Inches(max_width_inches))
        else:
            p.add_run().add_picture(str(image_bytes_or_path), width=Inches(max_width_inches))
    except Exception:
        run = p.add_run("(error inserting image)")
        set_run_style(run, size_pt=9, italic=True)


# ----------------------- Validation helper -----------------------
def validate_numeric_fields(form: Dict[str, Any], rules: List[Tuple[str, bool]]) -> List[Dict[str, str]]:
    """
    rules: list of tuples (field_name, is_integer)
      - is_integer True -> try int()
      - else -> try float()
    Returns list of error dicts: [{"field": name, "msg": "..."}]
    """
    errors = []
    for field, is_int in rules:
        val = form.get(field)
        if val is None or str(val).strip() == "":
            # empty allowed (we use placeholder)
            continue
        s = str(val).strip()
        try:
            if is_int:
                int(s)
            else:
                float(s)
        except Exception:
            errors.append({
                "field": field,
                "msg": f"Expected {'integer' if is_int else 'number'} but got '{s}'"
            })
    return errors


# ----------------------- Document builder functions (PLACEHOLDERS) -----------------------
# NOTE:
# Paste your implementations of the non-photo builder functions here (only once):
#
# - add_formulir_tallysheet(...)
# - add_elevasi_lahan_row(...)
# - add_kondisi_air_tanah_row(...)
# - add_tutupan_lahan_row(...)
# - add_flora_fauna_row(...)
# - add_drainase_row(...)
# - add_kualitas_air_row(...)
# - add_substratum_tanah_liat_row(...)
# - add_tipe_luapan_row(...)
# - add_ketebalan_gambut_row(...)
# - add_substratum_bawah_gambut_row(...)
# - add_perkembangan_kerusakan_row(...)
# - add_informasi_kebakaran_row(...)
# - add_analisis_lab_header(...)
# - add_porositas_row(...)
# - add_kelengasan_row(...)
# - add_c_organik_row(...)
# - add_sketsa_lokasi_row(...)
#
# (If you don't have them yet, tell me and I'll paste full implementations here.)
# ----------------------- End placeholders -----------------------


# ----------------------- Photo / "B. FOTO LAPANGAN" builders -----------------------
def add_foto_lapangan_section(
    doc: Document,
    images: Dict[str, Optional[Union[bytes, str]]] = None,
    max_image_width_inches: float = 2.8
):
    """
    Add section B. FOTO LAPANGAN (items 1..4).
    images keys (optional):
      - 'foto_air_tanah_genangan_1', 'foto_air_tanah_genangan_2'
      - 'foto_tutupan_lahan_1', 'foto_tutupan_lahan_2'
      - 'foto_flora_fauna_1', 'foto_flora_fauna_2'
      - 'foto_drainase_alami', 'foto_drainase_buatan'
    Values: bytes or path string.
    """
    if images is None:
        images = {}

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(4.5)

    # Header merged
    hdr = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    p = hdr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("B.\tFOTO  LAPANGAN")
    set_run_style(run, size_pt=12, bold=True)
    p_note = hdr.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = p_note.add_run("Seluruh hasil foto yang diambil harus jelas dan tidak membelakangi matahari")
    set_run_style(note_run, size_pt=9, italic=True)

    # Item 1
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("1.  Air tanah, genangan atau banjir")
    set_run_style(run, size_pt=11, bold=True)

    img_row = table.add_row()
    left, right = img_row.cells[0], img_row.cells[1]
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _insert_image_in_cell(left, images.get("foto_air_tanah_genangan_1"), max_image_width_inches)
    _insert_image_in_cell(right, images.get("foto_air_tanah_genangan_2"), max_image_width_inches)

    # Item 2
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("2.  Tutupan lahan, penggunaan lahan dan kondisinya")
    set_run_style(run, size_pt=11, bold=True)

    img_row = table.add_row()
    left, right = img_row.cells[0], img_row.cells[1]
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _insert_image_in_cell(left, images.get("foto_tutupan_lahan_1"), max_image_width_inches)
    _insert_image_in_cell(right, images.get("foto_tutupan_lahan_2"), max_image_width_inches)

    # Item 3
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("3.  Keberadaan flora dan fauna yang dilindungi")
    set_run_style(run, size_pt=11, bold=True)

    img_row = table.add_row()
    left, right = img_row.cells[0], img_row.cells[1]
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _insert_image_in_cell(left, images.get("foto_flora_fauna_1"), max_image_width_inches)
    _insert_image_in_cell(right, images.get("foto_flora_fauna_2"), max_image_width_inches)

    # Item 4 (drainase)
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("4.  Kondisi drainase alami dan drainase buatan")
    set_run_style(run, size_pt=11, bold=True)

    # subheader row
    sub_row = table.add_row()
    left, right = sub_row.cells[0], sub_row.cells[1]
    pleft = left.paragraphs[0]; pleft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pright = right.paragraphs[0]; pright.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rleft = pleft.add_run("Drainase alami"); set_run_style(rleft, size_pt=11, bold=True)
    rright = pright.add_run("Drainase buatan"); set_run_style(rright, size_pt=11, bold=True)

    img_row = table.add_row()
    left, right = img_row.cells[0], img_row.cells[1]
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _insert_image_in_cell(left, images.get("foto_drainase_alami"), max_image_width_inches)
    _insert_image_in_cell(right, images.get("foto_drainase_buatan"), max_image_width_inches)


def add_additional_photo_sections(
    doc: Document,
    images: Dict[str, Optional[Union[bytes, str]]] = None,
    max_image_width_inches: float = 2.8
):
    """
    Add photo sections 5..7 and 8..10 and 11..12 (all remaining photo blocks).
    images keys:
      - 'foto_kualitas_air_ec','foto_kualitas_air_tds','foto_kualitas_air_ph'
      - 'foto_tmat_1','foto_tmat_2'
      - 'foto_ketebalan_gambut_1','foto_ketebalan_gambut_2'
      - 'foto_substratum_ec','foto_substratum_ph'
      - 'foto_kerusakan_lahan_gambut_1','foto_kerusakan_lahan_gambut_2'
      - 'foto_karakteristik_tanah_pirit_1','foto_karakteristik_tanah_pirit_2'
      - 'foto_porositas_kelengasan_1','foto_porositas_kelengasan_2'
      - 'foto_tambahan_1' ... 'foto_tambahan_8'
    """
    if images is None:
        images = {}

    # Section 5: Kualitas Air / Kondisi Air Kanal (EC / TDS / pH)
    p5 = doc.add_paragraph(); p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r5 = p5.add_run("5.  Kualitas Air/Kondisi Air Kanal"); set_run_style(r5, size_pt=11, bold=True)

    tbl5 = doc.add_table(rows=2, cols=3)
    tbl5.autofit = False
    tbl5.columns[0].width = Inches(3.0)
    tbl5.columns[1].width = Inches(3.0)
    tbl5.columns[2].width = Inches(3.0)
    for i, lbl in enumerate(["EC", "TDS", "pH"]):
        c = tbl5.rows[0].cells[i]
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(lbl); set_run_style(run, size_pt=11, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # image row
    _insert_image_in_cell(tbl5.rows[1].cells[0], images.get("foto_kualitas_air_ec"), max_image_width_inches)
    _insert_image_in_cell(tbl5.rows[1].cells[1], images.get("foto_kualitas_air_tds"), max_image_width_inches)
    _insert_image_in_cell(tbl5.rows[1].cells[2], images.get("foto_kualitas_air_ph"), max_image_width_inches)

    doc.add_paragraph()

    # Section 6: TMAT (two images)
    p6 = doc.add_paragraph(); p6.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r6 = p6.add_run("6.  Pengukuran Tinggi Muka Air Tanah (TMAT) pada lubang bor titik pengamatan"); set_run_style(r6, size_pt=11, bold=True)

    tbl6 = doc.add_table(rows=1, cols=2); tbl6.autofit = False
    tbl6.columns[0].width = Inches(4.5); tbl6.columns[1].width = Inches(4.5)
    _insert_image_in_cell(tbl6.rows[0].cells[0], images.get("foto_tmat_1"), max_image_width_inches)
    _insert_image_in_cell(tbl6.rows[0].cells[1], images.get("foto_tmat_2"), max_image_width_inches)

    doc.add_paragraph()

    # Section 7: Ketebalan gambut (two images)
    p7 = doc.add_paragraph(); p7.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r7 = p7.add_run("7.  Ketebalan gambut"); set_run_style(r7, size_pt=11, bold=True)
    tbl7 = doc.add_table(rows=1, cols=2); tbl7.autofit = False
    tbl7.columns[0].width = Inches(4.5); tbl7.columns[1].width = Inches(4.5)
    _insert_image_in_cell(tbl7.rows[0].cells[0], images.get("foto_ketebalan_gambut_1"), max_image_width_inches)
    _insert_image_in_cell(tbl7.rows[0].cells[1], images.get("foto_ketebalan_gambut_2"), max_image_width_inches)

    doc.add_paragraph()

    # Section 8: Karakteristik substratum (EC / pH)
    p8 = doc.add_paragraph(); p8.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r8 = p8.add_run("8.  Karakteristik substratum dibawah lapisan gambut"); set_run_style(r8, size_pt=11, bold=True)
    tbl8 = doc.add_table(rows=2, cols=2); tbl8.autofit = False
    tbl8.columns[0].width = Inches(4.5); tbl8.columns[1].width = Inches(4.5)
    c_ec = tbl8.rows[0].cells[0].paragraphs[0]; c_ec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_ph = tbl8.rows[0].cells[1].paragraphs[0]; c_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_ec.add_run("EC"); set_run_style(c_ec.runs[0], size_pt=11, bold=True)
    c_ph.add_run("pH"); set_run_style(c_ph.runs[0], size_pt=11, bold=True)
    _insert_image_in_cell(tbl8.rows[1].cells[0], images.get("foto_substratum_ec"), max_image_width_inches)
    _insert_image_in_cell(tbl8.rows[1].cells[1], images.get("foto_substratum_ph"), max_image_width_inches)

    doc.add_paragraph()

    # Section 9: Perkembangan kerusakan (two images)
    p9 = doc.add_paragraph(); p9.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r9 = p9.add_run("9.  Perkembangan kondisi atau tingkat kerusakan lahan gambut (fungsi lindung/fungsi budidaya)"); set_run_style(r9, size_pt=11, bold=True)
    tbl9 = doc.add_table(rows=1, cols=2); tbl9.autofit = False
    tbl9.columns[0].width = Inches(4.5); tbl9.columns[1].width = Inches(4.5)
    _insert_image_in_cell(tbl9.rows[0].cells[0], images.get("foto_kerusakan_lahan_gambut_1"), max_image_width_inches)
    _insert_image_in_cell(tbl9.rows[0].cells[1], images.get("foto_kerusakan_lahan_gambut_2"), max_image_width_inches)

    doc.add_paragraph()

    # Section 10: Karakteristik tanah pirit (two images)
    p10 = doc.add_paragraph(); p10.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r10 = p10.add_run("10.  Karakteristik tanah dan kedalaman lapisan pirit"); set_run_style(r10, size_pt=11, bold=True)
    tbl10 = doc.add_table(rows=1, cols=2); tbl10.autofit = False
    tbl10.columns[0].width = Inches(4.5); tbl10.columns[1].width = Inches(4.5)
    _insert_image_in_cell(tbl10.rows[0].cells[0], images.get("foto_karakteristik_tanah_pirit_1"), max_image_width_inches)
    _insert_image_in_cell(tbl10.rows[0].cells[1], images.get("foto_karakteristik_tanah_pirit_2"), max_image_width_inches)

    doc.add_paragraph()

    # Section 11: Porositas & Kelengasan (two images)
    p11 = doc.add_paragraph(); p11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r11 = p11.add_run("11.  Porositas dan Kelengasan"); set_run_style(r11, size_pt=11, bold=True)
    tbl11 = doc.add_table(rows=1, cols=2); tbl11.autofit = False
    tbl11.columns[0].width = Inches(4.5); tbl11.columns[1].width = Inches(4.5)
    _insert_image_in_cell(tbl11.rows[0].cells[0], images.get("foto_porositas_kelengasan_1"), max_image_width_inches)
    _insert_image_in_cell(tbl11.rows[0].cells[1], images.get("foto_porositas_kelengasan_2"), max_image_width_inches)

    doc.add_paragraph()

    # Section 12: Foto Tambahan (up to 8 images arranged in 2 columns)
    p12 = doc.add_paragraph(); p12.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r12 = p12.add_run("12.  Foto Tambahan"); set_run_style(r12, size_pt=11, bold=True)

    tambahan_keys = [f"foto_tambahan_{i}" for i in range(1, 9)]
    pair_iter = [tambahan_keys[i:i+2] for i in range(0, len(tambahan_keys), 2)]
    for pair in pair_iter:
        tbl = doc.add_table(rows=1, cols=2); tbl.autofit = False
        tbl.columns[0].width = Inches(4.5); tbl.columns[1].width = Inches(4.5)
        left_key = pair[0]
        right_key = pair[1] if len(pair) > 1 else None
        _insert_image_in_cell(tbl.rows[0].cells[0], images.get(left_key), max_image_width_inches)
        if right_key:
            _insert_image_in_cell(tbl.rows[0].cells[1], images.get(right_key), max_image_width_inches)
        else:
            _insert_image_in_cell(tbl.rows[0].cells[1], None, max_image_width_inches)
        doc.add_paragraph()


# ----------------------- Endpoint: generate-full-section -----------------------
@app.post("/generate-full-section")
async def generate_full_section(
    background_tasks: BackgroundTasks,
    # koordinat (1)
    latitude_derajat: str = Form(None),
    latitude_menit: str = Form(None),
    latitude_detik: str = Form(None),
    latitude_arah: str = Form(None),
    longitude_derajat: str = Form(None),
    longitude_menit: str = Form(None),
    longitude_detik: str = Form(None),
    longitude_arah: str = Form(None),
    # elevasi (2)
    elevasi_lahan: str = Form(None),
    # kondisi air tanah (3)
    kedalaman_air_tanah: str = Form(None),
    genangan: str = Form(None),
    banjir_bulan: str = Form(None),
    banjir_lama_hari: str = Form(None),
    banjir_ketinggian_air: str = Form(None),
    sumber_air_hujan: str = Form(None),
    sumber_air_limpasan_sungai: str = Form(None),
    sumber_air_kiriman_hulu: str = Form(None),
    sumber_air_lainnya_checkbox: str = Form(None),
    sumber_air_lainnya_text: str = Form(None),
    # tutupan lahan (4)
    jenis_tanaman: str = Form(None),
    status_masyarakat: str = Form(None),
    status_perusahaan: str = Form(None),
    nama_perusahaan: str = Form(None),
    luas_konsesi: str = Form(None),
    # flora/fauna (5)
    flora_tidak_ada: str = Form(None),
    flora_ada: str = Form(None),
    flora_jenis: str = Form(None),
    fauna_tidak_ada: str = Form(None),
    fauna_ada: str = Form(None),
    fauna_jenis: str = Form(None),
    # drainase (6)
    drainase_alami_tidak_ada: str = Form(None),
    drainase_alami_ada: str = Form(None),
    drainase_buatan_tidak_ada: str = Form(None),
    drainase_buatan_ada: str = Form(None),
    drainase_buatan_saluran_terbuka: str = Form(None),
    drainase_buatan_saluran_terkontrol: str = Form(None),
    tinggi_muka_air_saluran: str = Form(None),
    # kualitas air (7)
    kualitas_air_tanah_ph: str = Form(None),
    kualitas_air_saluran_ph: str = Form(None),
    kualitas_air_tanah_ec: str = Form(None),
    kualitas_air_saluran_ec: str = Form(None),
    kualitas_air_tanah_tds: str = Form(None),
    kualitas_air_saluran_tds: str = Form(None),
    # substratum tanah liat (8)
    substratum_tanah_liat_ph: str = Form(None),
    substratum_tanah_liat_ec: str = Form(None),
    # tipe luapan (9)
    tipe_luapan_kemarau_a: str = Form(None),
    tipe_luapan_kemarau_b: str = Form(None),
    tipe_luapan_kemarau_c: str = Form(None),
    tipe_luapan_kemarau_d: str = Form(None),
    tipe_luapan_hujan_a: str = Form(None),
    tipe_luapan_hujan_b: str = Form(None),
    tipe_luapan_hujan_c: str = Form(None),
    tipe_luapan_hujan_d: str = Form(None),
    # ketebalan gambut (10)
    ketebalan_gambut_cm: str = Form(None),
    tingkat_perombakan_saprik: str = Form(None),
    tingkat_perombakan_hemik: str = Form(None),
    tingkat_perombakan_fibrik: str = Form(None),
    # substratum bawah (11)
    substratum_pasir_kwarsa: str = Form(None),
    substratum_clay_sedimen_sungai: str = Form(None),
    substratum_sedimen_berpirit: str = Form(None),
    substratum_granit: str = Form(None),
    substratum_lainnya_checkbox: str = Form(None),
    substratum_lainnya_text: str = Form(None),
    # perkembangan kerusakan (12)
    kerusakan_drainase_buatan: str = Form(None),
    kerusakan_terekspos_sedimen: str = Form(None),
    kondisi_tanaman_tidak_normal: str = Form(None),
    kondisi_tanaman_tidak_produktif: str = Form(None),
    kondisi_tanaman_miring_tumbang: str = Form(None),
    kondisi_tanaman_terjadi_subsiden_checkbox: str = Form(None),
    kondisi_tanaman_subsiden_cm: str = Form(None),
    kerapatan_tajuk: str = Form(None),
    # informasi kebakaran & hujan (13)
    kebakaran_tahun: str = Form(None),
    kebakaran_bulan: str = Form(None),
    kebakaran_tanggal: str = Form(None),
    kebakaran_lama_kejadian_bulan: str = Form(None),
    pemadaman_swadaya_masyarakat: str = Form(None),
    pemadaman_bantuan_pemerintah: str = Form(None),
    hujan_tanggal: str = Form(None),
    hujan_bulan: str = Form(None),
    hujan_tahun: str = Form(None),
    hujan_lama_kejadian_jam: str = Form(None),
    intensitas_hujan_tinggi: str = Form(None),
    intensitas_hujan_sedang: str = Form(None),
    intensitas_hujan_rendah: str = Form(None),
    # porositas / kelengasan / c-organik (14..16)
    porositas_bobot_isi: str = Form(None),
    kelengasan_kadar_air: str = Form(None),
    c_organik: str = Form(None),
    # sketsa lokasi (file)
    sketsa_lokasi_image: UploadFile = File(None),
    # photos: multiple optional UploadFile fields (example names)
    foto_air_tanah_genangan_1: UploadFile = File(None),
    foto_air_tanah_genangan_2: UploadFile = File(None),
    foto_tutupan_lahan_1: UploadFile = File(None),
    foto_tutupan_lahan_2: UploadFile = File(None),
    foto_flora_fauna_1: UploadFile = File(None),
    foto_flora_fauna_2: UploadFile = File(None),
    foto_drainase_alami: UploadFile = File(None),
    foto_drainase_buatan: UploadFile = File(None),
    foto_kualitas_air_ec: UploadFile = File(None),
    foto_kualitas_air_tds: UploadFile = File(None),
    foto_kualitas_air_ph: UploadFile = File(None),
    foto_tmat_1: UploadFile = File(None),
    foto_tmat_2: UploadFile = File(None),
    foto_ketebalan_gambut_1: UploadFile = File(None),
    foto_ketebalan_gambut_2: UploadFile = File(None),
    foto_substratum_ec: UploadFile = File(None),
    foto_substratum_ph: UploadFile = File(None),
    foto_kerusakan_lahan_gambut_1: UploadFile = File(None),
    foto_kerusakan_lahan_gambut_2: UploadFile = File(None),
    foto_karakteristik_tanah_pirit_1: UploadFile = File(None),
    foto_karakteristik_tanah_pirit_2: UploadFile = File(None),
    foto_porositas_kelengasan_1: UploadFile = File(None),
    foto_porositas_kelengasan_2: UploadFile = File(None),
    # foto tambahan 1..8
    foto_tambahan_1: UploadFile = File(None),
    foto_tambahan_2: UploadFile = File(None),
    foto_tambahan_3: UploadFile = File(None),
    foto_tambahan_4: UploadFile = File(None),
    foto_tambahan_5: UploadFile = File(None),
    foto_tambahan_6: UploadFile = File(None),
    foto_tambahan_7: UploadFile = File(None),
    foto_tambahan_8: UploadFile = File(None),
):
    # Collect form into dict for validation convenience
    form = dict(locals())

    # Build numeric validation rules (field_name, is_integer)
    numeric_rules = [
        ("latitude_derajat", False),
        ("latitude_menit", False),
        ("latitude_detik", False),
        ("longitude_derajat", False),
        ("longitude_menit", False),
        ("longitude_detik", False),
        ("elevasi_lahan", False),
        ("kedalaman_air_tanah", False),
        ("genangan", False),
        ("banjir_lama_hari", True),
        ("banjir_ketinggian_air", False),
        ("tinggi_muka_air_saluran", False),
        ("kualitas_air_tanah_ph", False),
        ("kualitas_air_saluran_ph", False),
        ("kualitas_air_tanah_ec", False),
        ("kualitas_air_saluran_ec", False),
        ("kualitas_air_tanah_tds", False),
        ("kualitas_air_saluran_tds", False),
        ("substratum_tanah_liat_ph", False),
        ("substratum_tanah_liat_ec", False),
        ("ketebalan_gambut_cm", False),
        ("kondisi_tanaman_subsiden_cm", False),
        ("hujan_lama_kejadian_jam", False),
        ("porositas_bobot_isi", False),
        ("kelengasan_kadar_air", False),
        ("c_organik", False),
    ]
    errors = validate_numeric_fields(form, numeric_rules)
    if errors:
        return JSONResponse(status_code=422, content={"detail": errors})

    # Build document
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.25)

    # --- Ensure add_* builder functions are implemented and available above ---
    add_formulir_tallysheet(
        doc,
        nomor="1",
        latitude_derajat=latitude_derajat,
        latitude_menit=latitude_menit,
        latitude_detik=latitude_detik,
        latitude_arah=latitude_arah,
        longitude_derajat=longitude_derajat,
        longitude_menit=longitude_menit,
        longitude_detik=longitude_detik,
        longitude_arah=longitude_arah,
    )

    add_elevasi_lahan_row(doc, nomor="2", elevasi_lahan=elevasi_lahan)

    add_kondisi_air_tanah_row(
        doc,
        kedalaman_air_tanah=kedalaman_air_tanah,
        genangan=genangan,
        banjir_bulan=banjir_bulan,
        banjir_lama_hari=banjir_lama_hari,
        banjir_ketinggian_air=banjir_ketinggian_air,
        sumber_air_hujan=sumber_air_hujan,
        sumber_air_limpasan_sungai=sumber_air_limpasan_sungai,
        sumber_air_kiriman_hulu=sumber_air_kiriman_hulu,
        sumber_air_lainnya_checkbox=sumber_air_lainnya_checkbox,
        sumber_air_lainnya_text=sumber_air_lainnya_text,
    )

    add_tutupan_lahan_row(
        doc,
        nomor="4",
        jenis_tanaman=jenis_tanaman,
        status_masyarakat=status_masyarakat,
        status_perusahaan=status_perusahaan,
        nama_perusahaan=nama_perusahaan,
        luas_konsesi=luas_konsesi,
    )

    add_flora_fauna_row(
        doc,
        numero="5" if False else "5",  # keep signature same; some earlier code used nomor param
        flora_tidak_ada=flora_tidak_ada,
        flora_ada=flora_ada,
        flora_jenis=flora_jenis,
        fauna_tidak_ada=fauna_tidak_ada,
        fauna_ada=fauna_ada,
        fauna_jenis=fauna_jenis,
    )

    add_drainase_row(
        doc,
        nomor="6",
        drainase_alami_tidak_ada=drainase_alami_tidak_ada,
        drainase_alami_ada=drainase_alami_ada,
        drainase_buatan_tidak_ada=drainase_buatan_tidak_ada,
        drainase_buatan_ada=drainase_buatan_ada,
        drainase_buatan_saluran_terbuka=drainase_buatan_saluran_terbuka,
        drainase_buatan_saluran_terkontrol=drainase_buatan_saluran_terkontrol,
        tinggi_muka_air_saluran=tinggi_muka_air_saluran,
    )

    # Section 7 .. 16 (assumes add_kualitas_air_row etc. are defined)
    add_kualitas_air_row(
        doc,
        nomor="7",
        kualitas_air_tanah_ph=kualitas_air_tanah_ph,
        kualitas_air_saluran_ph=kualitas_air_saluran_ph,
        kualitas_air_tanah_ec=kualitas_air_tanah_ec,
        kualitas_air_saluran_ec=kualitas_air_saluran_ec,
        kualitas_air_tanah_tds=kualitas_air_tanah_tds,
        kualitas_air_saluran_tds=kualitas_air_saluran_tds,
    )

    add_substratum_tanah_liat_row(
        doc,
        nomor="8",
        substratum_tanah_liat_ph=substratum_tanah_liat_ph,
        substratum_tanah_liat_ec=substratum_tanah_liat_ec,
    )

    add_tipe_luapan_row(
        doc,
        nomor="9",
        tipe_luapan_kemarau_a=tipe_luapan_kemarau_a,
        tipe_luapan_kemarau_b=tipe_luapan_kemarau_b,
        tipe_luapan_kemarau_c=tipe_luapan_kemarau_c,
        tipe_luapan_kemarau_d=tipe_luapan_kemarau_d,
        tipe_luapan_hujan_a=tipe_luapan_hujan_a,
        tipe_luapan_hujan_b=tipe_luapan_hujan_b,
        tipe_luapan_hujan_c=tipe_luapan_hujan_c,
        tipe_luapan_hujan_d=tipe_luapan_hujan_d,
    )

    add_ketebalan_gambut_row(
        doc,
        nomor="10",
        ketebalan_gambut_cm=ketebalan_gambut_cm,
        tingkat_perombakan_saprik=tingkat_perombakan_saprik,
        tingkat_perombakan_hemik=tingkat_perombakan_hemik,
        tingkat_perombakan_fibrik=tingkat_perombakan_fibrik,
    )

    add_substratum_bawah_gambut_row(
        doc,
        nomor="11",
        substratum_pasir_kwarsa=substratum_pasir_kwarsa,
        substratum_clay_sedimen_sungai=substratum_clay_sedimen_sungai,
        substratum_sedimen_berpirit=substratum_sedimen_berpirit,
        substratum_granit=substratum_granit,
        substratum_lainnya_checkbox=substratum_lainnya_checkbox,
        substratum_lainnya_text=substratum_lainnya_text,
    )

    add_perkembangan_kerusakan_row(
        doc,
        nomor="12",
        kerusakan_drainase_buatan=kerusakan_drainase_buatan,
        kerusakan_terekspos_sedimen=kerusakan_terekspos_sedimen,
        kondisi_tanaman_tidak_normal=kondisi_tanaman_tidak_normal,
        kondisi_tanaman_tidak_produktif=kondisi_tanaman_tidak_produktif,
        kondisi_tanaman_miring_tumbang=kondisi_tanaman_miring_tumbang,
        kondisi_tanaman_terjadi_subsiden_checkbox=kondisi_tanaman_terjadi_subsiden_checkbox,
        kondisi_tanaman_subsiden_cm=kondisi_tanaman_subsiden_cm,
        kerapatan_tajuk=kerapatan_tajuk,
    )

    add_informasi_kebakaran_row(
        doc,
        nomor="13",
        kebakaran_tahun=kebakaran_tahun,
        kebakaran_bulan=kebakaran_bulan,
        kebakaran_tanggal=kebakaran_tanggal,
        kebakaran_lama_kejadian_bulan=kebakaran_lama_kejadian_bulan,
        pemadaman_swadaya_masyarakat=pemadaman_swadaya_masyarakat,
        pemadaman_bantuan_pemerintah=pemadaman_bantuan_pemerintah,
        hujan_tanggal=hujan_tanggal,
        hujan_bulan=hujan_bulan,
        hujan_tahun=hujan_tahun,
        hujan_lama_kejadian_jam=hujan_lama_kejadian_jam,
        intensitas_hujan_tinggi=intensitas_hujan_tinggi,
        intensitas_hujan_sedang=intensitas_hujan_sedang,
        intensitas_hujan_rendah=intensitas_hujan_rendah,
    )

    add_analisis_lab_header(doc)

    add_porositas_row(doc, nomor="14", porositas_bobot_isi=porositas_bobot_isi)
    add_kelengasan_row(doc, nomor="15", kelengasan_kadar_air=kelengasan_kadar_air)
    add_c_organik_row(doc, nomor="16", c_organik=c_organik)

    # sketsa lokasi
    image_bytes = None
    if sketsa_lokasi_image is not None:
        try:
            image_bytes = await sketsa_lokasi_image.read()
        except Exception:
            image_bytes = None
    add_sketsa_lokasi_row(doc, sketsa_image_bytes=image_bytes)

    # collect uploaded photos to dict (read bytes)
    photo_fields = [
        "foto_air_tanah_genangan_1","foto_air_tanah_genangan_2",
        "foto_tutupan_lahan_1","foto_tutupan_lahan_2",
        "foto_flora_fauna_1","foto_flora_fauna_2",
        "foto_drainase_alami","foto_drainase_buatan",
        "foto_kualitas_air_ec","foto_kualitas_air_tds","foto_kualitas_air_ph",
        "foto_tmat_1","foto_tmat_2",
        "foto_ketebalan_gambut_1","foto_ketebalan_gambut_2",
        "foto_substratum_ec","foto_substratum_ph",
        "foto_kerusakan_lahan_gambut_1","foto_kerusakan_lahan_gambut_2",
        "foto_karakteristik_tanah_pirit_1","foto_karakteristik_tanah_pirit_2",
        "foto_porositas_kelengasan_1","foto_porositas_kelengasan_2",
        *[f"foto_tambahan_{i}" for i in range(1,9)]
    ]
    images: Dict[str, Optional[bytes]] = {}
    for fname in photo_fields:
        upload = form.get(fname)
        if hasattr(upload, "read"):
            try:
                images[fname] = await upload.read()
            except Exception:
                images[fname] = None
        else:
            images[fname] = None

    # add photo sections
    add_foto_lapangan_section(doc, images=images)
    add_additional_photo_sections(doc, images=images)

    # Save to temp file and schedule cleanup
    tmpf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmpf_name = tmpf.name
    tmpf.close()
    doc.save(tmpf_name)

    out_name = f"tallysheet_{uuid.uuid4().hex}.docx"

    def _cleanup(path: str):
        try:
            os.remove(path)
        except Exception:
            pass

    background_tasks.add_task(_cleanup, tmpf_name)

    return FileResponse(
        path=tmpf_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=out_name,
    )


# ----------------------- Endpoint: generate-sample -----------------------
@app.get("/generate-sample")
def generate_sample(background_tasks: BackgroundTasks):
    """
    Create a sample tallysheet docx (filled with example values) to inspect layout.
    """
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.25)

    # Example usage (requires add_* functions defined)
    add_formulir_tallysheet(
        doc,
        nomor="1",
        latitude_derajat="1",
        latitude_menit="23",
        latitude_detik="45",
        latitude_arah="LS",
        longitude_derajat="102",
        longitude_menit="10",
        longitude_detik="12",
        longitude_arah="BT",
    )
    add_elevasi_lahan_row(doc, nomor="2", elevasi_lahan="12")
    add_kondisi_air_tanah_row(
        doc,
        kedalaman_air_tanah="45",
        genangan="10",
        banjir_bulan="Januari",
        banjir_lama_hari="5",
        banjir_ketinggian_air="20",
        sumber_air_hujan="on",
        sumber_air_limpasan_sungai="",
        sumber_air_kiriman_hulu="on",
        sumber_air_lainnya_checkbox="",
        sumber_air_lainnya_text="drainase tersumbat",
    )
    add_tutupan_lahan_row(
        doc,
        nomor="4",
        jenis_tanaman="Kelapa",
        status_masyarakat="on",
        status_perusahaan="",
        nama_perusahaan="",
        luas_konsesi="",
    )
    add_flora_fauna_row(doc, numero="5" if False else "5", flora_tidak_ada="", flora_ada="on", flora_jenis="Rafflesia", fauna_tidak_ada="", fauna_ada="on", fauna_jenis="Orangutan")
    add_drainase_row(doc, nomor="6", drainase_alami_tidak_ada="", drainase_alami_ada="on", drainase_buatan_tidak_ada="", drainase_buatan_ada="on", drainase_buatan_saluran_terbuka="on", drainase_buatan_saluran_terkontrol="", tinggi_muka_air_saluran="30")
    add_kualitas_air_row(doc, nomor="7", kualitas_air_tanah_ph="4.2", kualitas_air_saluran_ph="5.6", kualitas_air_tanah_ec="120", kualitas_air_saluran_ec="200", kualitas_air_tanah_tds="50", kualitas_air_saluran_tds="80")
    add_substratum_tanah_liat_row(doc, nomor="8", substratum_tanah_liat_ph="6.6", substratum_tanah_liat_ec="300")
    add_tipe_luapan_row(doc, nomor="9", tipe_luapan_kemarau_a="on", tipe_luapan_kemarau_b="", tipe_luapan_kemarau_c="", tipe_luapan_kemarau_d="", tipe_luapan_hujan_a="", tipe_luapan_hujan_b="on", tipe_luapan_hujan_c="", tipe_luapan_hujan_d="")
    add_ketebalan_gambut_row(doc, nomor="10", ketebalan_gambut_cm="120", tingkat_perombakan_saprik="on", tingkat_perombakan_hemik="", tingkat_perombakan_fibrik="")
    add_substratum_bawah_gambut_row(doc, nomor="11", substratum_pasir_kwarsa="on", substratum_clay_sedimen_sungai="", substratum_sedimen_berpirit="", substratum_granit="", substratum_lainnya_checkbox="", substratum_lainnya_text="")
    add_perkembangan_kerusakan_row(doc, nomor="12", kerusakan_drainase_buatan="on", kerusakan_terekspos_sedimen="", kondisi_tanaman_tidak_normal="on", kondisi_tanaman_tidak_produktif="", kondisi_tanaman_miring_tumbang="", kondisi_tanaman_terjadi_subsiden_checkbox="", kondisi_tanaman_subsiden_cm="", kerapatan_tajuk="3")
    add_informasi_kebakaran_row(doc, nomor="13", kebakaran_tahun="2023", kebakaran_bulan="Agustus", kebakaran_tanggal="15", kebakaran_lama_kejadian_bulan="1", pemadaman_swadaya_masyarakat="on", pemadaman_bantuan_pemerintah="", hujan_tanggal="10", hujan_bulan="07", hujan_tahun="2024", hujan_lama_kejadian_jam="6", intensitas_hujan_tinggi="on", intensitas_hujan_sedang="", intensitas_hujan_rendah="")
    add_analisis_lab_header(doc)
    add_porositas_row(doc, nomor="14", porositas_bobot_isi="45")
    add_kelengasan_row(doc, nomor="15", kelengasan_kadar_air="85")
    add_c_organik_row(doc, nomor="16", c_organik="58")
    add_sketsa_lokasi_row(doc, sketsa_image_bytes=None)

    # sample photos: pass None so placeholders appear
    add_foto_lapangan_section(doc, images={})
    add_additional_photo_sections(doc, images={})

    tmpf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmpf_name = tmpf.name
    tmpf.close()
    doc.save(tmpf_name)

    out_name = f"tallysheet_sample_{uuid.uuid4().hex}.docx"

    def _cleanup(path: str):
        try:
            os.remove(path)
        except Exception:
            pass

    background_tasks.add_task(_cleanup, tmpf_name)

    return FileResponse(
        path=tmpf_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=out_name,
    )